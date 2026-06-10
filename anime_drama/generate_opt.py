# -*- coding: utf-8 -*-
"""第1集优化版 - 分镜衔接更自然 + 减少群演 + 节奏更好"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
doc = Document()
s = doc.styles['Normal']; s.font.name = 'Microsoft YaHei'; s.font.size = Pt(11)

t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("第1集：叶青云·玉佩异变（优化版）")
r.font.size = Pt(22); r.bold = True
t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run("改编自《修仙：我叶青云，绝顶之上》第一章  |  约75秒  |  9:16竖屏")
r2.font.size = Pt(12); r2.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph("")
doc.add_heading("角色表", level=1)
for n, d in [
    ("叶青云","16岁，黑发垂眼，灰色旧袍，眼神沉郁隐忍，身型偏瘦，但脊背挺直"),
    ("叶川","嫡系二少爷，青蓝锦袍，面容高傲，眼神轻蔑，身后跟着两个跟班"),
    ("跟班甲","瘦高个，贼眉鼠眼，跟在叶川身后谄笑"),
    ("跟班乙","矮壮，叉腰站着，一脸鄙夷"),
    ("母亲","病弱苍白，躺在床上，面容温柔"),
]:
    p = doc.add_paragraph(); r = p.add_run("【%s】"%n); r.bold = True; p.add_run("  "+d)

doc.add_page_break()

scenes = [
    ("祭祖大典·被辱（12s）", [
        "【场景】叶家大院，祭祖香火缭绕。远景：族人背影肃立，只有两三个侧面剪影",
        "【人物】叶青云、叶川、跟班甲、跟班乙",
        "(叶青云站在人群末尾，垂眼低头，手不自觉摸向怀中玉佩)",
        "叶川(高声讥讽，走近两步)：废物就是废物，连三叔家的狗都比你强！",
        "(跟班甲在旁嗤笑，跟班乙叉腰撇嘴)",
        "叶川：听说你炼气二层卡了三年了？你爹跑了，你娘死了——趁早滚出叶家，省得给我们丢人！",
        "(叶青云缓缓抬头，眼神漆黑如深渊，手指攥紧了玉佩)",
        "叶青云(低沉但清晰)：你说我可以——说我爹不行。",
    ]),
    ("玉佩·闪回（8s）", [
        "【场景】从叶青云攥紧玉佩的拳头特写开始——叠化到七年前昏暗的病榻",
        "【人物】叶青云(幼年)、母亲",
        "(特写：粗糙的拳头中露出一截黑色玉佩，温润微光)",
        "(叠化：七年前，母亲枯瘦的手将同一枚玉佩放入小叶青云掌心)",
        "母亲(气若游丝)：这是你爹留给你的……好好保管……",
        "(叠化回到现实：叶青云站在院中，指节发白，眼中有泪光但没落下)",
    ]),
    ("偏院·深夜苦修（10s）", [
        "【场景】破旧厢房，屋顶漏雨，月光斜射入屋，空荡荡只有一张破床",
        "【人物】叶青云",
        "(叶青云盘坐在床上，月光照在他脸上，闭目运转引气诀)",
        "(灵气刚在丹田聚起一丝，就像被什么吞噬了，消散无踪)",
        "(他猛地睁眼，一拳砸在床板上，灰尘簌簌落下)",
        "叶青云(低哑)：又失败了……三日后……去做杂役？",
        "(他低头看着自己的双手，指节因握拳而发白)",
    ]),
    ("觉醒·玉佩异变（20s）", [
        "【场景】偏院中，月光皎洁。窗外偶尔传来虫鸣，然后突然安静",
        "【人物】叶青云",
        "(叶青云烦躁站起身踱步，手不经意触到怀里的玉佩——)",
        "(指尖猛地一缩！像被烫到一样！)",
        "(特写：黑色玉佩在月色下泛着幽幽暗金色光芒，微微颤动)",
        "叶青云(惊退一步)：怎么回事？！",
        "(玉佩像活了一样，一道暗金色流光从玉佩飞出，直冲他的眉心——)",
        "(轰——！)",
        "(黑暗混沌空间浮现，中央悬浮一枚暗金色光球，沉重古老的气息扑面而来)",
        "(《混沌天经》——四个大字如雷霆般炸响！)",
        "(外界，夜空风云变色！方圆十里灵气汇聚成巨大白色旋涡，从偏院上方灌入！)",
        "(远景：叶家正堂方向，隐约可见有人影冲出观望，但偏院才是焦点)",
    ]),
    ("突破·结尾悬念（10s）", [
        "【场景】偏院屋内，晨光从破窗洒入。灵气已经平息",
        "【人物】叶青云",
        "(叶青云猛地睁眼——瞳孔中一道暗金色光芒一闪而逝)",
        "(缓缓抬起右手，皮肤表面覆盖着淡淡金色光泽，经脉中灵力奔涌)",
        "(握拳，指节咔咔作响，嘴角缓缓勾起——六年了)",
        "字幕：三年的废物……从今夜起，不再是了。",
        "(画面渐黑，字幕浮现)",
        "第2集预告：青石碑前，一拳惊人。",
    ]),
]

for idx, (title, lines) in enumerate(scenes, 1):
    doc.add_heading("第%d场 %s" % (idx, title), level=1)
    for line in lines:
        if line.startswith("【"):
            p = doc.add_paragraph(); r = p.add_run(line); r.italic = True; r.font.color.rgb = RGBColor(80,80,80)
        elif line.startswith("("):
            p = doc.add_paragraph(); r = p.add_run(line); r.italic = True; r.font.color.rgb = RGBColor(120,120,120)
        elif "：" in line:
            parts = line.split("：",1)
            p = doc.add_paragraph(); r = p.add_run(parts[0]+"："); r.bold = True; p.add_run(parts[1] if len(parts)>1 else "")
        else:
            doc.add_paragraph(line)
    doc.add_paragraph("")

out = os.path.join(ROOT, "anime_drama", "episode_01", "episode_01_opt.docx")
doc.save(out)
print("OK: " + out)
